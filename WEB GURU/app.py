from flask import Flask, render_template, url_for, request, redirect, jsonify, send_file, session, flash
from werkzeug.utils import secure_filename
from io import BytesIO
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from docx import Document
from models import db, Guru, User, LogTable
from werkzeug.security import generate_password_hash
from werkzeug.security import check_password_hash
from datetime import datetime
from werkzeug.utils import secure_filename
from decorators import login_required
import os
import pytz
from functools import wraps


app = Flask(__name__)

#============KONFIGURASI DATABASE============#
app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+mysqlconnector://root:@localhost/sistem_pengolahan_data_guru'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.secret_key = 'adsi'

# Inisialisasi db dengan Flask
db.init_app(app)

#============UPLOAD FOTO===========#
# Tentukan folder untuk menyimpan upload
app.config['UPLOAD_FOLDER'] = 'static/uploads'  # Tentukan folder upload
app.config['ALLOWED_EXTENSIONS'] = {'png', 'jpg', 'jpeg', 'gif'}  # Format file yang diperbolehkan

# Fungsi untuk memeriksa ekstensi file yang valid
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

#==========LOG TABLE===========#
@app.route('/log')
def log():
    logs = LogTable.query.order_by(LogTable.timestamp.desc()).all()  # Ambil semua log dari tabel
    for log in logs:
        # Mengonversi waktu UTC ke WIB dan format sebelum dikirim ke template
        log.timestamp = log.timestamp.astimezone(pytz.timezone('Asia/Jakarta'))
    return render_template('log.html', logs=logs)

@app.route('/reset_log', methods=['POST'])
@login_required
def reset_log():
    try:
        # Menghapus semua data log
        db.session.query(LogTable).delete()
        db.session.commit()
        flash("Semua log telah dihapus!", "success")
    except Exception as e:
        flash(f"Terjadi kesalahan: {str(e)}", "danger")
    
    return redirect(url_for('log'))

def log_activity(action, table_name, record_id, log_message=None):
    try:
        # Menambahkan log ke tabel log_table
        new_log = LogTable(
            action=action,
            table_name=table_name,
            record_id=record_id,
            log_message=log_message,
            timestamp=datetime.utcnow()
        )
        db.session.add(new_log)
        db.session.commit()
        print("Log berhasil ditambahkan!")
    except Exception as e:
        print(f"Terjadi kesalahan saat menambahkan log: {str(e)}")

#==========USER,LOGIN,LOGOUT,REGISTER==========#
@app.route('/user')
def user_dashboard():
    user_id = session.get('user_id')  # Ambil ID pengguna dari sesi
    user = User.query.get(user_id)  # Ambil data pengguna dari database
    return render_template('user.html', username=user.username)  # Kirimkan username ke template

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        # Cek apakah username ada di database
        user = User.query.filter_by(username=username).first()

        if user and check_password_hash(user.password, password):
            session['user_id'] = user.id  # Menyimpan ID pengguna di session
            
            # Tambahkan log untuk login
            new_log = LogTable(
                action="login",
                table_name="user",
                record_id=user.id,
                log_message=f"User {user.username} logged in."
            )
            new_log.set_timestamp()
            db.session.add(new_log)
            db.session.commit()

            flash("Login successful!", "success")
            return redirect(url_for('index'))  # Arahkan ke halaman utama setelah login
        else:
            flash("Invalid username or password", "danger")
    
    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    # Cek apakah pengguna sudah login
    user_logged_in = None
    if 'user_id' in session:
        user_logged_in = User.query.get(session['user_id'])
    
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        # Hash password sebelum menyimpannya
        hashed_password = generate_password_hash(password)
        
        # Simpan pengguna baru ke database
        new_user = User(username=username, password=hashed_password)
        db.session.add(new_user)
        db.session.commit()

        # Tambahkan log untuk register
        new_log = LogTable(
            action="register",
            table_name="user",
            record_id=new_user.id,
            log_message=f"User {new_user.username} registered."
        )
        new_log.set_timestamp()
        db.session.add(new_log)
        db.session.commit()

        flash("Account successfully created!", "success")
        return redirect(url_for('login'))
    
    return render_template('register.html', user_logged_in=user_logged_in)

@app.route('/logout')
def logout():
    user_id = session.get('user_id')  # Ambil ID pengguna dari sesi
    if user_id:
        user = User.query.get(user_id)

        # Tambahkan log untuk logout
        if user:
            new_log = LogTable(
                action="logout",
                table_name="user",
                record_id=user.id,
                log_message=f"User {user.username} logged out."
            )
            new_log.set_timestamp()
            db.session.add(new_log)
            db.session.commit()

    session.clear()  # Menghapus sesi login
    flash('You have been logged out.', 'success')
    return redirect(url_for('login'))


#=============HALAMAN UTAMA=============#
@app.route('/')
@login_required
def index():
    data_guru = Guru.query.all()
    return render_template('index.html', data_guru=data_guru)

#=============GURU DETAIL=============#
@app.route('/guru-detail/<int:id_guru>', methods=['GET'])
def guru_detail(id_guru):
    guru = Guru.query.get(id_guru)
    if guru:
        return jsonify({
            'nama_guru': guru.nama_guru,
            'nip': guru.nip,
            'alamat': guru.alamat,
            'no_telepon': guru.no_telepon,
            'mata_pelajaran': guru.mata_pelajaran,
            'foto': guru.foto or 'default_image_url.jpg'  # Ganti dengan URL gambar default jika foto tidak ada
        })
    return jsonify({'error': 'Guru not found'}), 404

#============CREATE GURU============#
@app.route('/create-guru', methods=['GET', 'POST'])
def create_guru():
    if request.method == 'POST':
        # Mengambil data dari form
        nama_guru = request.form['nama_guru']
        nip = request.form['nip'] if request.form.get('status_guru') != 'honorer' else None
        alamat = request.form['alamat']
        no_telepon = request.form['no_telepon']
        mata_pelajaran = request.form['mata_pelajaran']
        status_guru = request.form['status_guru']  # Mengambil status guru dari form
        
        # Menangani upload foto
        foto = request.files['foto']
        if foto:
            foto_filename = secure_filename(foto.filename)
            foto.save(os.path.join(app.config['UPLOAD_FOLDER'], foto_filename))
        else:
            foto_filename = None
        
        # Menyimpan data guru ke database
        guru = Guru(
            nama_guru=nama_guru,
            nip=nip if status_guru != "honorer" else None, 
            alamat=alamat,
            no_telepon=no_telepon,
            mata_pelajaran=mata_pelajaran,
            status_guru=status_guru,  # Menambahkan status_guru
            foto=foto_filename
        )
        db.session.add(guru)
        db.session.commit()

        return redirect(url_for('index'))  # Arahkan ke halaman utama setelah data disimpan

    return render_template('create_guru.html')

#=================EDIT GURU=================#
@app.route('/edit-guru/<int:id_guru>', methods=['GET', 'POST'])
def edit_guru(id_guru):
    guru = Guru.query.get_or_404(id_guru)

    if request.method == 'POST':
        guru.nama_guru = request.form['nama_guru']
        guru.nip = request.form['nip']
        guru.alamat = request.form['alamat']
        guru.no_telepon = request.form['no_telepon']
        guru.mata_pelajaran = request.form['mata_pelajaran']

        # Menangani upload foto baru
        foto_baru = request.files.get('foto')
        if foto_baru:
            guru.foto = foto_baru.filename
            foto_baru.save(f'./static/uploads/{guru.foto}')

        # Menyimpan perubahan ke database
        db.session.commit()

        return redirect('/')

    return render_template('edit_guru.html', guru=guru)

#==================DELETE GURU============#
@app.route('/delete-guru/<int:id_guru>', methods=['GET'])
def delete_guru(id_guru):
    guru = Guru.query.get_or_404(id_guru)

    # Menghapus foto jika ada
    if guru.foto:
        os.remove(f'./static/uploads/{guru.foto}')

    # Menghapus data guru dari database
    db.session.delete(guru)
    db.session.commit()

    return redirect('/')

#===============PEMBUATAN LAPORAN============#
@app.route('/buat-laporan', methods=['GET', 'POST'])
def buat_laporan():
    # Mengambil semua data guru dari database
    gurus = Guru.query.all()

    # Jika metode POST, akan menghasilkan laporan dalam format yang diinginkan
    if request.method == 'POST':
        laporan_format = request.form['format']

        if laporan_format == 'pdf':
            return generate_pdf(gurus)
        elif laporan_format == 'excel':
            return generate_excel(gurus)
        elif laporan_format == 'word':
            return generate_word(gurus)

    return render_template('buat_laporan.html', gurus=gurus)

#=============FILE PDF===========#
def generate_pdf(gurus):
    # Membuat file PDF di memori
    output = BytesIO()
    c = canvas.Canvas(output, pagesize=letter)
    width, height = letter  # Dapatkan ukuran halaman

    # Judul laporan
    c.setFont("Helvetica-Bold", 14)
    c.drawString(100, height - 50, "Laporan Data Guru")
    c.setFont("Helvetica", 10)
    
    # Header tabel
    c.drawString(100, height - 80, "Nama Guru")
    c.drawString(300, height - 80, "NIP")
    c.drawString(500, height - 80, "Alamat")
    c.line(100, height - 82, 500, height - 82)  # Membuat garis bawah header

    y_position = height - 100  # Posisi mulai untuk data guru

    # Menulis data guru dalam bentuk tabel
    for guru in gurus:
        # Menangani NIP jika None
        nip_value = guru.nip if guru.nip is not None else 'Tidak Ada NIP'
        
        c.drawString(100, y_position, guru.nama_guru)  # Menulis Nama Guru
        c.drawString(300, y_position, nip_value)       # Menulis NIP, gunakan nip_value
        c.drawString(500, y_position, guru.alamat)     # Menulis Alamat
        y_position -= 20  # Posisi baris berikutnya
        
        if y_position < 100:  # Jika halaman sudah penuh, buat halaman baru
            c.showPage()
            y_position = height - 50  # Reset posisi y untuk halaman baru
            c.setFont("Helvetica-Bold", 14)
            c.drawString(100, height - 50, "Laporan Data Guru")
            c.setFont("Helvetica", 10)
            c.drawString(100, height - 80, "Nama Guru")
            c.drawString(300, height - 80, "NIP")
            c.drawString(500, height - 80, "Alamat")
            c.line(100, height - 82, 500, height - 82)

    c.save()

    # Mengembalikan file PDF sebagai respons download
    output.seek(0)
    return send_file(output, as_attachment=True, download_name="laporan_guru.pdf", mimetype='application/pdf')

#============FILE EXCEL==========#
def generate_excel(gurus):
    # Mengambil data guru dan mengonversinya ke DataFrame
    data = []
    for guru in gurus:
        data.append([guru.nama_guru, guru.nip, guru.alamat, guru.no_telepon, guru.mata_pelajaran])

    df = pd.DataFrame(data, columns=["Nama Guru", "NIP", "Alamat", "No Telepon", "Mata Pelajaran"])

    # Menyimpan ke dalam file Excel
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Data Guru')

    output.seek(0)
    return send_file(output, as_attachment=True, download_name="laporan_guru.xlsx", mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

#===========FILE WORD=========#
def generate_word(gurus):
    # Membuat dokumen Word
    doc = Document()
    doc.add_heading('Laporan Data Guru', 0)

    # Menambahkan tabel dengan data guru
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nama Guru'
    hdr_cells[1].text = 'NIP'
    hdr_cells[2].text = 'Alamat'
    hdr_cells[3].text = 'No Telepon'
    hdr_cells[4].text = 'Mata Pelajaran'

    # Menambahkan data guru ke tabel
    for guru in gurus:
        row_cells = table.add_row().cells
        row_cells[0].text = guru.nama_guru
        
        # Menangani NIP jika None
        nip_value = guru.nip if guru.nip is not None else 'Tidak Ada NIP'
        row_cells[1].text = nip_value  # Menulis NIP atau 'Tidak Ada NIP'
        
        row_cells[2].text = guru.alamat
        row_cells[3].text = guru.no_telepon
        row_cells[4].text = guru.mata_pelajaran

    # Menyimpan dokumen Word
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return send_file(output, as_attachment=True, download_name="laporan_guru.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == "__main__":
    with app.app_context():
        db.create_all()  # Membuat tabel di database
    app.run(debug=True)